from docx import Document
import os
import glob
import time
import re
from education import models as education_models
from health import models as health_models


def media_url_full(request, media_file_path):
    """
    Return the full URL of the media file
    """
    return f"{request.scheme}://{request.META['HTTP_HOST']}{media_file_path}"


# Used in clean_html function, but compiled here once for performance improvements
CLEANR = re.compile('<.*?>')


def clean_html(raw_html):
    """
    Remove all tags from HTML and convert chars
    """
    clean_text = re.sub(CLEANR, '', raw_html)
    clean_text = clean_text\
        .replace('&#39;', "'")\
        .replace('&nbsp;', ' ')\
        .replace('&quot;', '"')
    return clean_text


def create_document(request):
    """
    Creates a Word Document (.docx) and returns its file path
    """

    # Delete all existing files in the data folder
    data_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'data')
    files = glob.glob(data_path + '/*')
    for f in files:
        os.remove(f)

    # Establish new file name
    file_name = f'encv_data_{time.strftime("%Y-%m-%d_%H-%M")}.docx'
    file_path = os.path.join(data_path, file_name)

    # Create a new Document
    document = Document()

    # A visual separator for data items presented in the document
    item_separator = """

----------------------------------------------------------------------------------------------------------------------

"""

    # Build content within Document:
    # Introduction content
    document.add_heading('ENCV Data Export', 0)
    document.add_paragraph("""
This document contains data exported from the ENCV database.

Data is organised into 2 strands:
1) Education Strand (includes a list of 'Journal Entries' from participants)
2) Health Strand (includes 'Conversations' between cancer champions and the patients)
""")
    document.add_page_break()
    # Education strand content
    document.add_heading('1) Education Strand', 1)
    document.add_paragraph(item_separator)
    for journal_entry in education_models.JournalEntry.objects.all().select_related('author',).prefetch_related('prompt').order_by('id'):
        document.add_heading(f'Journal Entry ID: {journal_entry.id}', 2)
        document.add_paragraph(f"""
Author:
{str(journal_entry.author)}

Created:
{str(journal_entry.created)[:16]}

Last Updated:
{str(journal_entry.last_updated)[:16]}

Link:
{journal_entry.link}

Image (download link):
{media_url_full(request, journal_entry.image.url) if journal_entry.image else None}

Audio (download link):
{media_url_full(request, journal_entry.audio.url) if journal_entry.audio else None}

Video (download link):
{media_url_full(request, journal_entry.video.url) if journal_entry.video else None}

Prompt(s):
{journal_entry.prompts_as_str}

Journal Entry Text:
{clean_html(journal_entry.text)}

""")
        document.add_paragraph(item_separator)
    # Health strand content
    document.add_page_break()
    document.add_page_break()
    document.add_heading('2) Health Strand', 1)
    document.add_paragraph(item_separator)
    for conversation in health_models.Conversation.objects.all().select_related('author',).order_by('id'):
        document.add_heading(f'Conversation ID: {conversation.id}', 2)
        document.add_paragraph(f"""
Author:
{str(conversation.author)}

Created:
{str(conversation.created)[:16]}

Last Updated:
{str(conversation.last_updated)[:16]}

Conversation Date:
{str(conversation.conversation_date)[:10]}

Conversation Audio (download link):
{media_url_full(request, conversation.conversation_audio.url) if conversation.conversation_audio else None}

Conversation Transcript (download link):
{media_url_full(request, conversation.conversation_transcript.url) if conversation.conversation_transcript else None}

Cancer Champion Reflection:
{conversation.cancer_champion_reflection}

""")
        document.add_paragraph(item_separator)

    # Save document and return the file path
    document.save(file_path)
    return file_path
